import tkinter as tk
from tkinter import messagebox
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Função para criar o documento docx formatado
def create_vpn_doc(username, password):
    # Cria o documento
    doc = Document()

    # Adiciona uma tabela com duas linhas e duas colunas
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'

    # Formatação do texto da tabela
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(12)

    # Linha 1: Nome do usuário
    table.cell(0, 0).text = "Nome:"
    table.cell(0, 1).text = username

    # Linha 2: Senha
    table.cell(1, 0).text = "Senha:"
    table.cell(1, 1).text = password

    # Salva o arquivo
    doc_name = "Usuario_Senha_Details.docx"
    doc.save(doc_name)
    messagebox.showinfo("Sucesso", f"Documento salvo como {doc_name}")

# Função chamada quando o botão é clicado
def generate_doc():
    username = entry_user.get().strip()
    password = entry_pass.get().strip()

    if not username or not password:
        messagebox.showerror("Erro", "Por favor, insira o nome do usuário e a senha.")
    else:
        create_vpn_doc(username, password)

# Interface gráfica usando Tkinter
root = tk.Tk()
root.title("Gerador de Documento com Usuário e Senha")
root.geometry("400x200")

# Label e entrada para o nome do usuário
label_user = tk.Label(root, text="Usuário:")
label_user.pack(pady=10)
entry_user = tk.Entry(root, width=40)
entry_user.pack()

# Label e entrada para a senha
label_pass = tk.Label(root, text="Senha:")
label_pass.pack(pady=10)
entry_pass = tk.Entry(root, width=40)
entry_pass.pack()

# Botão para gerar o documento
generate_button = tk.Button(root, text="Gerar Documento", command=generate_doc)
generate_button.pack(pady=20)

# Inicia a interface gráfica
root.mainloop()
