import tkinter as tk
from tkinter import messagebox
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# Função para criar o documento docx formatado
def create_vpn_doc(users_passwords):
    # Cria o documento
    doc = Document()

    for i, (user, password) in enumerate(users_passwords):
        # Adiciona a tabela com uma linha para o título e mais linhas para aplicação, usuário e senha
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'

        # Linha 1: Título "OPEN VPN" ocupando as duas colunas
        title_cell = table.cell(0, 0)
        title_cell.merge(table.cell(0, 1))  # Mescla as duas células para o título
        title_paragraph = title_cell.paragraphs[0]
        run = title_paragraph.add_run("OPEN VPN")
        run.font.name = 'Roboto'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 100, 175)  # Cor #0064af
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centraliza o título

        # Linha 2: Aplicação
        cell1 = table.cell(1, 0)
        run1 = cell1.paragraphs[0].add_run('Aplicação')
        run1.font.name = 'Roboto'
        run1.font.size = Pt(10)
        run1.font.bold = True

        cell2 = table.cell(1, 1)
        run2 = cell2.paragraphs[0].add_run('Open VPN')
        run2.font.name = 'Calibri'
        run2.font.size = Pt(11)

        # Linha 3: Usuário
        cell3 = table.cell(2, 0)
        run3 = cell3.paragraphs[0].add_run('Usuário')
        run3.font.name = 'Roboto'
        run3.font.size = Pt(10)
        run3.font.bold = True

        cell4 = table.cell(2, 1)
        run4 = cell4.paragraphs[0].add_run(user)
        run4.font.name = 'Calibri'
        run4.font.size = Pt(11)

        # Linha 4: Senha
        cell5 = table.cell(3, 0)
        run5 = cell5.paragraphs[0].add_run('Senha')
        run5.font.name = 'Roboto'
        run5.font.size = Pt(10)
        run5.font.bold = True

        cell6 = table.cell(3, 1)
        run6 = cell6.paragraphs[0].add_run(password if password else "(sem senha)")
        run6.font.name = 'Calibri'
        run6.font.size = Pt(11)

        # Adiciona uma linha vazia entre cada bloco de usuário e senha
        doc.add_paragraph("")

    # Salva o arquivo
    doc_name = "Usuarios_Senhas_Details.docx"
    doc.save(doc_name)
    messagebox.showinfo("Sucesso", f"Documento salvo como {doc_name}")

# Função chamada quando o botão é clicado
def generate_doc():
    users_passwords_text = text_input.get("1.0", tk.END).strip()

    if not users_passwords_text:
        messagebox.showerror("Erro", "Por favor, insira os usuários e senhas.")
        return

    users_passwords = []
    current_user = None
    current_password = None

    lines = users_passwords_text.splitlines()

    for line in lines:
        line = line.strip()

        if not line:  # Se houver uma linha vazia, pula
            continue

        if current_user is None:
            current_user = line  # Considera a primeira linha não vazia como o nome do usuário
        elif current_password is None:
            current_password = line  # Considera a segunda linha como a senha

        if current_user and current_password:
            users_passwords.append((current_user, current_password))
            current_user = None  # Reseta para o próximo par
            current_password = None

    # Se sobrar um último usuário sem senha, adiciona ao final
    if current_user is not None:
        users_passwords.append((current_user, "(sem senha)"))

    create_vpn_doc(users_passwords)

# Interface gráfica usando Tkinter
root = tk.Tk()
root.title("Gerador de Documento com Usuário e Senha")
root.geometry("600x400")

# Label para instrução
label_instruction = tk.Label(root, text="Insira usuários e senhas (um por linha):")
label_instruction.pack(pady=10)

# Caixa de texto para inserir os usuários e senhas
text_input = tk.Text(root, height=10, width=50)
text_input.pack(pady=10)

# Botão para gerar o documento
generate_button = tk.Button(root, text="Gerar Documento", command=generate_doc)
generate_button.pack(pady=20)

# Inicia a interface gráfica
root.mainloop()
