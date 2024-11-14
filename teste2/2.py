from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_vpn_doc(application, user, password):
    # Cria o documento
    doc = Document()

    # Adiciona o título "OPEN VPN"
    title = doc.add_heading('OPEN VPN', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Adiciona a tabela
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'

    # Adiciona os dados na tabela
    # Linha 1: Aplicação
    cell1 = table.cell(0, 0)
    cell1.text = 'Aplicação'
    cell1.paragraphs[0].runs[0].font.bold = True

    cell2 = table.cell(0, 1)
    cell2.text = application

    # Linha 2: Usuário
    cell3 = table.cell(1, 0)
    cell3.text = 'Usuário'
    cell3.paragraphs[0].runs[0].font.bold = True

    cell4 = table.cell(1, 1)
    cell4.text = user

    # Linha 3: Senha
    cell5 = table.cell(2, 0)
    cell5.text = 'Senha'
    cell5.paragraphs[0].runs[0].font.bold = True

    cell6 = table.cell(2, 1)
    cell6.text = password

    # Salva o arquivo
    doc_name = "Open_VPN_Details.docx"
    doc.save(doc_name)
    print(f"Documento salvo como {doc_name}")

# Exemplo de uso
application_name = input("Digite o nome da aplicação: ")
username = input("Digite o nome do usuário: ")
password = input("Digite a senha: ")

create_vpn_doc(application_name, username, password)
