import tkinter as tk
from tkinter import messagebox
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Função para criar o documento docx formatado
def create_vpn_doc(users_passwords):
    # Cria o documento
    doc = Document()

    for i, (user, password) in enumerate(users_passwords):
        # Adiciona o título "USUÁRIO E SENHA" para cada entrada (pode ser adaptado para cada bloco se necessário)
        title = doc.add_heading(f'Usuário {i + 1}', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Adiciona a tabela
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'

        # Adiciona os dados na tabela
        # Linha 1: Usuário
        cell1 = table.cell(0, 0)
        cell1.text = 'Usuário'
        cell1.paragraphs[0].runs[0].font.bold = True

        cell2 = table.cell(0, 1)
        cell2.text = user

        # Linha 2: Senha
        cell3 = table.cell(1, 0)
        cell3.text = 'Senha'
        cell3.paragraphs[0].runs[0].font.bold = True

        cell4 = table.cell(1, 1)
        cell4.text = password

        # Adiciona uma linha vazia entre cada bloco de usuário e senha
        doc.add_paragraph("")

    # Salva o arquivo
    doc_name = "Usuarios_Senhas_Details.docx"
    doc.save(doc_name)
    messagebox.showinfo("Sucesso", f"Documento salvo como {doc_name}")

# Função chamada quando o botão é clicado
def generate_doc():
    users_passwords_text = text_input.get("1.0", tk.END).strip()

    if not users_passwords_text:
        messagebox.showerror("Erro", "Por favor, insira os usuários e senhas.")
        return

    users_passwords = []
    lines = users_passwords_text.splitlines()

    # Processar as linhas, assumindo que os pares são "usuário" e "senha"
    for i in range(0, len(lines), 2):
        if i + 1 < len(lines):
            user = lines[i].strip()
            password = lines[i + 1].strip()
            users_passwords.append((user, password))

    # Verifica se há um número ímpar de linhas e adiciona um aviso
    if len(lines) % 2 != 0:
        messagebox.showwarning("Aviso", "Há um número ímpar de entradas. A última entrada será ignorada se não houver senha correspondente.")

    create_vpn_doc(users_passwords)

# Interface gráfica usando Tkinter
root = tk.Tk()
root.title("Gerador de Documento com Usuário e Senha")
root.geometry("600x400")

# Label para instrução
label_instruction = tk.Label(root, text="Insira usuários e senhas (um por linha):")
label_instruction.pack(pady=10)

# Caixa de texto para inserir os usuários e senhas
text_input = tk.Text(root, height=10, width=50)
text_input.pack(pady=10)

# Botão para gerar o documento
generate_button = tk.Button(root, text="Gerar Documento", command=generate_doc)
generate_button.pack(pady=20)

# Inicia a interface gráfica
root.mainloop()
