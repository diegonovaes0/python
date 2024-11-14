import tkinter as tk
from tkinter import messagebox
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Função para definir o fundo da célula e aplicar bordas pretas de 1 pt
def set_cell_background_and_borders(cell, color, border_size):
    """Define o fundo branco e aplica bordas pretas de 1 pt na célula."""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    
    # Define o fundo branco
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

    # Define bordas pretas de 1 pt
    for border_type in ['top', 'left', 'bottom', 'right']:
        tc_borders = OxmlElement(f'w:{border_type}')
        tc_borders.set(qn('w:val'), 'single')
        tc_borders.set(qn('w:sz'), str(border_size))  # Tamanho da linha (1 pt = 8 em OOXML units)
        tc_borders.set(qn('w:color'), '000000')  # Cor preta
        tcPr.append(tc_borders)

# Função para aplicar indentação de 1 espaço no início de cada linha
def apply_indentation(paragraph):
    """Aplica uma indentação de 0,15 cm (aproximadamente 1 espaço) no início da linha."""
    paragraph.paragraph_format.first_line_indent = Cm(0.15)  # Aproximadamente 1 espaço

# Função para criar o documento docx formatado
def create_vpn_doc(users_passwords):
    # Cria o documento
    doc = Document()

    for i, (user, password) in enumerate(users_passwords):
        # Adiciona a tabela com uma linha para o título e mais linhas para aplicação, usuário e senha
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'

        # Linha 1: Título "OPEN VPN" ocupando as duas colunas
        title_cell = table.cell(0, 0)
        title_cell.merge(table.cell(0, 1))  # Mescla as duas células para o título
        title_paragraph = title_cell.paragraphs[0]
        run = title_paragraph.add_run("OPEN VPN")
        run.font.name = 'Roboto'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 100, 175)  # Cor #0064af
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centraliza o título
        set_cell_background_and_borders(title_cell, 'FFFFFF', 8)  # Fundo branco, bordas pretas 1 pt
        apply_indentation(title_paragraph)

        # Linha 2: Aplicação
        cell1 = table.cell(1, 0)
        run1 = cell1.paragraphs[0].add_run('Aplicação')
        run1.font.name = 'Roboto'
        run1.font.size = Pt(10)
        run1.font.bold = True
        set_cell_background_and_borders(cell1, 'FFFFFF', 8)  # Fundo branco, bordas pretas 1 pt
        apply_indentation(cell1.paragraphs[0])

        cell2 = table.cell(1, 1)
        run2 = cell2.paragraphs[0].add_run('Open VPN')
        run2.font.name = 'Calibri'
        run2.font.size = Pt(11)
        set_cell_background_and_borders(cell2, 'FFFFFF', 8)  # Fundo branco, bordas pretas 1 pt
        apply_indentation(cell2.paragraphs[0])

        # Linha 3: Usuário
        cell3 = table.cell(2, 0)
        run3 = cell3.paragraphs[0].add_run('Usuário')
        run3.font.name = 'Roboto'
        run3.font.size = Pt(10)
        run3.font.bold = True
        set_cell_background_and_borders(cell3, 'FFFFFF', 8)  # Fundo branco, bordas pretas 1 pt
        apply_indentation(cell3.paragraphs[0])

        cell4 = table.cell(2, 1)
        run4 = cell4.paragraphs[0].add_run(user)
        run4.font.name = 'Calibri'
        run4.font.size = Pt(11)
        set_cell_background_and_borders(cell4, 'FFFFFF', 8)  # Fundo branco, bordas pretas 1 pt
        apply_indentation(cell4.paragraphs[0])

        # Linha 4: Senha
        cell5 = table.cell(3, 0)
        run5 = cell5.paragraphs[0].add_run('Senha')
        run5.font.name = 'Roboto'
        run5.font.size = Pt(10)
        run5.font.bold = True
        set_cell_background_and_borders(cell5, 'FFFFFF', 8)  # Fundo branco, bordas pretas 1 pt
        apply_indentation(cell5.paragraphs[0])

        cell6 = table.cell(3, 1)
        run6 = cell6.paragraphs[0].add_run(password if password else "(sem senha)")
        run6.font.name = 'Calibri'
        run6.font.size = Pt(11)
        set_cell_background_and_borders(cell6, 'FFFFFF', 8)  # Fundo branco, bordas pretas 1 pt
        apply_indentation(cell6.paragraphs[0])

        # Adiciona uma linha vazia entre cada bloco de usuário e senha
        doc.add_paragraph("")

    # Salva o arquivo
    doc_name = "Usuarios_Senhas_Details.docx"
    doc.save(doc_name)
    messagebox.showinfo("Sucesso", f"Documento salvo como {doc_name}")

# Função chamada quando o botão é clicado
def generate_doc():
    users_passwords_text = text_input.get("1.0", tk.END).strip()

    if not users_passwords_text:
        messagebox.showerror("Erro", "Por favor, insira os usuários e senhas.")
        return

    users_passwords = []
    current_user = None
    current_password = None

    lines = users_passwords_text.splitlines()

    for line in lines:
        line = line.strip()

        if not line:  # Se houver uma linha vazia, pula
            continue

        if current_user is None:
            current_user = line  # Considera a primeira linha não vazia como o nome do usuário
        elif current_password is None:
            current_password = line  # Considera a segunda linha como a senha

        if current_user and current_password:
            users_passwords.append((current_user, current_password))
            current_user = None  # Reseta para o próximo par
            current_password = None

    # Se sobrar um último usuário sem senha, adiciona ao final
    if current_user is not None:
        users_passwords.append((current_user, "(sem senha)"))

    create_vpn_doc(users_passwords)

# Interface gráfica usando Tkinter
root = tk.Tk()
root.title("Gerador de Documento com Usuário e Senha")
root.geometry("600x400")

# Label para instrução
label_instruction = tk.Label(root, text="Insira usuários e senhas (um por linha):")
label_instruction.pack(pady=10)

# Caixa de texto para inserir os usuários e senhas
text_input = tk.Text(root, height=10, width=50)
text_input.pack(pady=10)

# Botão para gerar o documento
generate_button = tk.Button(root, text="Gerar Documento", command=generate_doc)
generate_button.pack(pady=20)

# Inicia a interface gráfica
root.mainloop()
