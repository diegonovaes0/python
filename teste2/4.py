import tkinter as tk
from tkinter import messagebox
from docx import Document
from docx.shared import Pt

# Função para criar o documento docx formatado
def create_vpn_doc(users_passwords):
    # Cria o documento
    doc = Document()

    for user, password in users_passwords:
        # Adiciona a linha do usuário e senha
        doc.add_paragraph(f"Usuário: {user}")
        doc.add_paragraph(f"Senha: {password}")
        doc.add_paragraph("")  # Adiciona uma linha vazia para separação

    # Salva o arquivo
    doc_name = "Usuarios_Senhas_Details.docx"
    doc.save(doc_name)
    messagebox.showinfo("Sucesso", f"Documento salvo como {doc_name}")

# Função chamada quando o botão é clicado
def generate_doc():
    users_passwords_text = text_input.get("1.0", tk.END).strip()

    if not users_passwords_text:
        messagebox.showerror("Erro", "Por favor, insira os usuários e senhas.")
        return

    users_passwords = []
    lines = users_passwords_text.splitlines()

    # Processar as linhas, assumindo que os pares são "usuário" e "senha"
    for i in range(0, len(lines), 2):
        if i + 1 < len(lines):
            user = lines[i].strip()
            password = lines[i + 1].strip()
            users_passwords.append((user, password))

    create_vpn_doc(users_passwords)

# Interface gráfica usando Tkinter
root = tk.Tk()
root.title("Gerador de Documento com Usuário e Senha")
root.geometry("600x400")

# Label para instrução
label_instruction = tk.Label(root, text="Insira usuários e senhas (um por linha):")
label_instruction.pack(pady=10)

# Caixa de texto para inserir os usuários e senhas
text_input = tk.Text(root, height=10, width=50)
text_input.pack(pady=10)

# Botão para gerar o documento
generate_button = tk.Button(root, text="Gerar Documento", command=generate_doc)
generate_button.pack(pady=20)

# Inicia a interface gráfica
root.mainloop()
