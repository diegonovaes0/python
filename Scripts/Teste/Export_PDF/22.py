import tkinter as tk
from tkinter import filedialog, simpledialog, messagebox
from tkinter import ttk
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches

def create_document():
    # Cria o documento
    return Document()

def add_header(doc, header_text):
    # Adiciona o cabeçalho
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = header_text
    header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def add_footer(doc, footer_text):
    # Adiciona o rodapé
    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.text = footer_text
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def add_paragraph(doc, text, font_size=12):
    # Adiciona um parágrafo com tamanho de fonte personalizado
    paragraph = doc.add_paragraph(text)
    run = paragraph.runs[0]
    run.font.size = Pt(font_size)

def add_image(doc, image_path):
    # Adiciona uma imagem ao documento
    doc.add_picture(image_path, width=Inches(4))

def save_document(doc):
    # Abre um diálogo para salvar o arquivo
    file_path = filedialog.asksaveasfilename(
        title="Salvar documento",
        defaultextension=".docx",
        filetypes=[("Documentos Word", "*.docx")]
    )
    if file_path:
        doc.save(file_path)
        messagebox.showinfo("Sucesso", f"Documento salvo em: {file_path}")

# Funções para manipular a interface
def add_header_ui():
    header_text = simpledialog.askstring("Cabeçalho", "Digite o texto do cabeçalho:")
    if header_text:
        add_header(doc, header_text)

def add_footer_ui():
    footer_text = simpledialog.askstring("Rodapé", "Digite o texto do rodapé:")
    if footer_text:
        add_footer(doc, footer_text)

def add_paragraph_ui():
    text = simpledialog.askstring("Parágrafo", "Digite o texto do parágrafo:")
    if text:
        font_size = simpledialog.askinteger("Fonte", "Digite o tamanho da fonte (em pontos):", initialvalue=12)
        add_paragraph(doc, text, font_size)

def add_image_ui():
    # Abre um diálogo para selecionar a imagem
    image_path = filedialog.askopenfilename(
        title="Selecionar Imagem",
        filetypes=[("Imagens", "*.png;*.jpg;*.jpeg;*.bmp;*.gif")]
    )
    if image_path:
        add_image(doc, image_path)

# Configuração da Interface Gráfica
doc = create_document()
root = tk.Tk()
root.title("Editor de Documentos DOCX")

# Botões para adicionar conteúdo
ttk.Button(root, text="Adicionar Cabeçalho", command=add_header_ui).pack(pady=5)
ttk.Button(root, text="Adicionar Rodapé", command=add_footer_ui).pack(pady=5)
ttk.Button(root, text="Adicionar Parágrafo", command=add_paragraph_ui).pack(pady=5)
ttk.Button(root, text="Adicionar Imagem", command=add_image_ui).pack(pady=5)
ttk.Button(root, text="Salvar Documento", command=lambda: save_document(doc)).pack(pady=20)

root.mainloop()
