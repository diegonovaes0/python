import json
from docx import Document
import tkinter as tk
from tkinter import filedialog, messagebox

def convert_docx_to_json(docx_path, output_json):
    # Abre o arquivo .docx
    doc = Document(docx_path)
    content = {"paragraphs": [], "tables": []}

    # Extrai parágrafos
    for paragraph in doc.paragraphs:
        content["paragraphs"].append(paragraph.text)

    # Extrai tabelas
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        content["tables"].append(table_data)

    # Salva o conteúdo em um arquivo JSON
    with open(output_json, 'w', encoding='utf-8') as json_file:
        json.dump(content, json_file, ensure_ascii=False, indent=4)

    messagebox.showinfo("Sucesso", f"Conteúdo do documento salvo em {output_json}")

def select_file_and_convert():
    # Abre a janela para selecionar o arquivo .docx
    file_path = filedialog.askopenfilename(
        title="Selecione o arquivo .docx",
        filetypes=[("Documentos Word", "*.docx")]
    )

    if file_path:
        # Define o caminho de saída para o JSON
        output_json = filedialog.asksaveasfilename(
            title="Salvar como JSON",
            defaultextension=".json",
            filetypes=[("Arquivos JSON", "*.json")]
        )

        if output_json:
            # Converte o documento para JSON
            convert_docx_to_json(file_path, output_json)

# Interface Gráfica
root = tk.Tk()
root.title("Conversor de DOCX para JSON")

# Botão para iniciar a conversão
button = tk.Button(root, text="Selecionar Arquivo .docx e Converter", command=select_file_and_convert)
button.pack(pady=20)

root.mainloop()
