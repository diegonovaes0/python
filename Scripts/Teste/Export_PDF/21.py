import tkinter as tk
from tkinter import filedialog, simpledialog, messagebox
from docx import Document

def load_docx(file_path):
    # Carrega o arquivo .docx
    return Document(file_path)

def edit_document(doc, placeholder, new_text):
    # Procura pelo placeholder e substitui pelo novo texto
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, new_text)
    
    # Edita o texto nas tabelas, se houver
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if placeholder in cell.text:
                    cell.text = cell.text.replace(placeholder, new_text)

def save_docx(doc, output_path):
    # Salva o documento editado
    doc.save(output_path)
    messagebox.showinfo("Sucesso", f"Documento salvo em: {output_path}")

def select_file_and_edit():
    # Seleciona o arquivo .docx
    file_path = filedialog.askopenfilename(
        title="Selecione o arquivo .docx",
        filetypes=[("Documentos Word", "*.docx")]
    )

    if file_path:
        doc = load_docx(file_path)

        # Solicita o texto a ser substituído e o novo texto
        placeholder = simpledialog.askstring("Entrada", "Digite o texto a ser substituído:")
        new_text = simpledialog.askstring("Entrada", "Digite o novo texto:")

        if placeholder and new_text:
            edit_document(doc, placeholder, new_text)

            # Seleciona o local para salvar o novo documento
            output_path = filedialog.asksaveasfilename(
                title="Salvar documento editado",
                defaultextension=".docx",
                filetypes=[("Documentos Word", "*.docx")]
            )

            if output_path:
                save_docx(doc, output_path)

# Interface Gráfica
root = tk.Tk()
root.title("Editor de Documentos DOCX")

# Botão para selecionar e editar o arquivo
button = tk.Button(root, text="Selecionar Arquivo .docx e Editar", command=select_file_and_edit)
button.pack(pady=20)

root.mainloop()
