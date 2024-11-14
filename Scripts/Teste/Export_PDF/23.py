import requests
from docx import Document
from io import BytesIO

def download_docx_from_google_drive(file_id):
    # URL de download direto do Google Drive
    download_url = f"https://drive.google.com/uc?id={file_id}&export=download"
    
    # Baixa o arquivo
    response = requests.get(download_url)
    if response.status_code == 200:
        print("Download concluído com sucesso.")
        return BytesIO(response.content)
    else:
        print("Falha ao baixar o arquivo.")
        return None

def read_docx_content(docx_file):
    # Carrega o documento do conteúdo baixado
    doc = Document(docx_file)
    content = {"paragraphs": [], "tables": []}

    # Extrai parágrafos
    for paragraph in doc.paragraphs:
        content["paragraphs"].append(paragraph.text)

    # Extrai tabelas
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        content["tables"].append(table_data)

    return content

# Insira o ID do arquivo aqui
file_id = 'SEU_DOCUMENTO_ID'  # Substitua pelo ID do seu documento no Google Drive

# Baixa e processa o documento
docx_file = download_docx_from_google_drive(file_id)
if docx_file:
    content = read_docx_content(docx_file)

    # Exibe o conteúdo extraído
    print("Parágrafos:")
    for paragraph in content["paragraphs"]:
        print(paragraph)

    print("\nTabelas:")
    for table in content["tables"]:
        for row in table:
            print(row)
