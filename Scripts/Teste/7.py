import tkinter as tk
from tkinter import messagebox
from docx import Document
from docx.shared import Pt, RGBColor
import random
import string

# Função para gerar senha forte
def generate_password(length=14):
    upper = random.choice(string.ascii_uppercase)
    lower = random.choice(string.ascii_lowercase)
    digit = random.choice(string.digits)
    special = random.choice(string.punctuation)
    
    remaining_length = length - 4
    all_characters = string.ascii_letters + string.digits + string.punctuation
    remaining_chars = ''.join(random.choice(all_characters) for _ in range(remaining_length))
    
    password = upper + lower + digit + special + remaining_chars
    password = ''.join(random.sample(password, len(password)))
    
    return password

# Função para gerar múltiplas senhas
def generate_multiple_passwords():
    num_passwords = len(users)
    password_length = int(password_length_spinbox.get())
    global passwords
    passwords = [generate_password(password_length) for _ in range(num_passwords)]
    passwords_text.delete(1.0, tk.END)
    passwords_text.insert(tk.END, "\n".join(passwords))
    merge_users_passwords(users)

# Função para mesclar automaticamente usuários e senhas
def merge_users_passwords(users_list):
    global passwords
    if len(users_list) != len(passwords):
        messagebox.showwarning("Erro", "O número de usuários e senhas não corresponde.")
        return
    
    merged_list = []
    for user, password in zip(users_list, passwords):
        merged_list.append(f"{user}\n{password}\n")

    merged_text.delete(1.0, tk.END)
    merged_text.insert(tk.END, "\n".join(merged_list))

# Função para criar o documento docx formatado
def create_vpn_doc(users_passwords):
    doc = Document()

    for i, (user, password) in enumerate(users_passwords):
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'

        # Linha 1: Usuário
        cell1 = table.cell(0, 0)
        run1 = cell1.paragraphs[0].add_run('Usuário')
        run1.font.name = 'Roboto'
        run1.font.size = Pt(12)
        run1.font.bold = True

        cell2 = table.cell(0, 1)
        run2 = cell2.paragraphs[0].add_run(user)
        run2.font.name = 'Calibri'
        run2.font.size = Pt(11)

        # Linha 2: Senha
        cell3 = table.cell(1, 0)
        run3 = cell3.paragraphs[0].add_run('Senha')
        run3.font.name = 'Roboto'
        run3.font.size = Pt(12)
        run3.font.bold = True

        cell4 = table.cell(1, 1)
        run4 = cell4.paragraphs[0].add_run(password if password else "(sem senha)")
        run4.font.name = 'Calibri'
        run4.font.size = Pt(11)

        doc.add_paragraph("")  # Adiciona uma linha vazia entre os usuários

    # Salva o arquivo
    doc_name = "Usuarios_Senhas_Details.docx"
    doc.save(doc_name)
    messagebox.showinfo("Sucesso", f"Documento salvo como {doc_name}")

# Função chamada quando o botão "Gerar Documento" é clicado
def generate_doc():
    users_passwords_text = merged_text.get("1.0", tk.END).strip()

    if not users_passwords_text:
        messagebox.showerror("Erro", "Por favor, insira os usuários e senhas.")
        return

    users_passwords = []
    current_user = None
    current_password = None

    lines = users_passwords_text.splitlines()

    for line in lines:
        line = line.strip()

        if not line:
            continue

        if current_user is None:
            current_user = line
        elif current_password is None:
            current_password = line

        if current_user and current_password:
            users_passwords.append((current_user, current_password))
            current_user = None
            current_password = None

    if current_user is not None:
        users_passwords.append((current_user, "(sem senha)"))

    create_vpn_doc(users_passwords)

# Interface gráfica
root = tk.Tk()
root.title("Gerador de Documento com Usuário e Senha")
root.geometry("750x600")

# Configuração da linha superior para o nome da empresa e quantidade de usuários
frame_top = tk.Frame(root)
frame_top.grid(row=0, column=0, columnspan=5, pady=10)

tk.Label(frame_top, text="Nome da Empresa:").grid(row=0, column=0, padx=(0, 5), sticky="e")
company_entry = tk.Entry(frame_top, width=20)
company_entry.grid(row=0, column=1, padx=(0, 20))

tk.Label(frame_top, text="Quantidade de Usuários:").grid(row=0, column=2, padx=(0, 5), sticky="e")
num_users_spinbox = tk.Spinbox(frame_top, from_=1, to=100, width=5)
num_users_spinbox.grid(row=0, column=3, padx=(0, 20))

create_users_button = tk.Button(frame_top, text="Criar Usuários", command=create_users)
create_users_button.grid(row=0, column=4, padx=(5, 0))

# Configuração da linha para comprimento da senha e botão de geração
frame_password = tk.Frame(root)
frame_password.grid(row=1, column=0, columnspan=5, pady=(0, 10))

tk.Label(frame_password, text="Comprimento da Senha:").grid(row=0, column=0, padx=(0, 5), sticky="e")
password_length_spinbox = tk.Spinbox(frame_password, from_=8, to=20, width=5)
password_length_spinbox.grid(row=0, column=1, padx=(0, 20))

generate_button = tk.Button(frame_password, text="Gerar Senhas", command=generate_multiple_passwords)
generate_button.grid(row=0, column=2, padx=(5, 0))

# Configuração para o campo de senhas geradas e o campo de usuários
frame_fields = tk.Frame(root)
frame_fields.grid(row=2, column=0, columnspan=5, pady=(0, 10))

tk.Label(frame_fields, text="Senhas Geradas").grid(row=0, column=0, padx=10, pady=(0, 5))
passwords_text = tk.Text(frame_fields, height=10, width=30)
passwords_text.grid(row=1, column=0, padx=10, pady=(0, 5))

tk.Label(frame_fields, text="Usuários").grid(row=0, column=1, padx=10, pady=(0, 5))
users_text = tk.Text(frame_fields, height=10, width=30)
users_text.grid(row=1, column=1, padx=10, pady=(0, 5))

# Legenda para o campo de usuários
tk.Label(frame_fields, text="Cole aqui os usuários para serem mesclados", fg="red", font=("TkDefaultFont", 8)).grid(row=2, column=1, padx=10, pady=(0, 10), sticky="n")

# Campo para exibir a mesclagem
tk.Label(root, text="Mesclagem de Usuários e Senhas").grid(row=3, column=0, columnspan=5, padx=10, pady=(10, 5))
merged_text = tk.Text(root, height=10, width=70)
merged_text.grid(row=4, column=0, columnspan=5, padx=10, pady=5)

# Botão para gerar o documento final
doc_button = tk.Button(root, text="Gerar Documento", command=generate_doc)
doc_button.grid(row=5, column=0, columnspan=5, pady=(10, 20))

users = []
passwords = []

root.mainloop()
