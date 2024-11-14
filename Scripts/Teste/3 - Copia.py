import tkinter as tk
from tkinter import messagebox
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import random
import string

# Função para definir o fundo da célula e aplicar bordas pretas de 1 pt
def set_cell_background_and_borders(cell, color, border_size):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

    for border_type in ['top', 'left', 'bottom', 'right']:
        tc_borders = OxmlElement(f'w:{border_type}')
        tc_borders.set(qn('w:val'), 'single')
        tc_borders.set(qn('w:sz'), str(border_size))
        tc_borders.set(qn('w:color'), '000000')
        tcPr.append(tc_borders)

# Função para criar o documento docx formatado
def create_vpn_doc(users_passwords):
    doc = Document()

    for i, (user, password) in enumerate(users_passwords):
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'

        title_cell = table.cell(0, 0)
        title_cell.merge(table.cell(0, 1))
        title_paragraph = title_cell.paragraphs[0]
        run = title_paragraph.add_run("OPEN VPN")
        run.font.name = 'Roboto'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 100, 175)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_background_and_borders(title_cell, 'FFFFFF', 8)

        cell1 = table.cell(1, 0)
        run1 = cell1.paragraphs[0].add_run('Aplicação')
        run1.font.name = 'Roboto'
        run1.font.size = Pt(10)
        run1.font.bold = True
        set_cell_background_and_borders(cell1, 'FFFFFF', 8)

        cell2 = table.cell(1, 1)
        run2 = cell2.paragraphs[0].add_run('Open VPN')
        run2.font.name = 'Calibri'
        run2.font.size = Pt(11)
        set_cell_background_and_borders(cell2, 'FFFFFF', 8)

        cell3 = table.cell(2, 0)
        run3 = cell3.paragraphs[0].add_run('Usuário')
        run3.font.name = 'Roboto'
        run3.font.size = Pt(10)
        run3.font.bold = True
        set_cell_background_and_borders(cell3, 'FFFFFF', 8)

        cell4 = table.cell(2, 1)
        run4 = cell4.paragraphs[0].add_run(user)
        run4.font.name = 'Calibri'
        run4.font.size = Pt(11)
        set_cell_background_and_borders(cell4, 'FFFFFF', 8)

        cell5 = table.cell(3, 0)
        run5 = cell5.paragraphs[0].add_run('Senha')
        run5.font.name = 'Roboto'
        run5.font.size = Pt(10)
        run5.font.bold = True
        set_cell_background_and_borders(cell5, 'FFFFFF', 8)

        cell6 = table.cell(3, 1)
        run6 = cell6.paragraphs[0].add_run(password if password else "(sem senha)")
        run6.font.name = 'Calibri'
        run6.font.size = Pt(11)
        set_cell_background_and_borders(cell6, 'FFFFFF', 8)

        doc.add_paragraph("")

    doc_name = "Usuarios_Senhas_Details.docx"
    doc.save(doc_name)
    messagebox.showinfo("Sucesso", f"Documento salvo como {doc_name}")

# Função para gerar senha forte
def generate_password(length=14):
    upper = random.choice(string.ascii_uppercase)
    lower = random.choice(string.ascii_lowercase)
    digit = random.choice(string.digits)
    special = random.choice(string.punctuation)
    
    remaining_length = length - 4
    all_characters = string.ascii_letters + string.digits + string.punctuation
    remaining_chars = ''.join(random.choice(all_characters) for _ in range(remaining_length))
    
    password = upper + lower + digit + special + remaining_chars
    password = ''.join(random.sample(password, len(password)))
    
    return password

# Função para gerar múltiplas senhas
def generate_multiple_passwords():
    num_passwords = len(users)
    password_length = int(password_length_spinbox.get())
    global passwords
    passwords = [generate_password(password_length) for _ in range(num_passwords)]
    passwords_text.delete(1.0, tk.END)
    passwords_text.insert(tk.END, "\n".join(passwords))
    merge_users_passwords(users)

# Função para mesclar automaticamente usuários e senhas
def merge_users_passwords(users_list):
    global passwords
    if len(users_list) != len(passwords):
        messagebox.showwarning("Erro", "O número de usuários e senhas não corresponde.")
        return
    
    merged_list = []
    for user, password in zip(users_list, passwords):
        merged_list.append(f"{user}\n{password}\n")

    merged_text.delete(1.0, tk.END)
    merged_text.insert(tk.END, "\n".join(merged_list))

# Função chamada quando o botão é clicado para gerar documento
def generate_doc():
    users_passwords_text = merged_text.get("1.0", tk.END).strip()

    if not users_passwords_text:
        messagebox.showerror("Erro", "Por favor, insira os usuários e senhas.")
        return

    users_passwords = []
    current_user = None
    current_password = None

    lines = users_passwords_text.splitlines()

    for line in lines:
        line = line.strip()

        if not line:
            continue

        if current_user is None:
            current_user = line
        elif current_password is None:
            current_password = line

        if current_user and current_password:
            users_passwords.append((current_user, current_password))
            current_user = None
            current_password = None

    if current_user is not None:
        users_passwords.append((current_user, "(sem senha)"))

    create_vpn_doc(users_passwords)

# Função para formatar os nomes dos usuários
def format_username(username):
    words = username.split()
    return ''.join(word.capitalize() for word in words)

# Função para atualizar a lista de usuários e acionar a mesclagem
def on_users_text_change(event):
    global users
    users = [format_username(user.strip()) for user in users_text.get(1.0, tk.END).strip().split('\n') if user.strip()]
    generate_multiple_passwords()

# Interface gráfica
root = tk.Tk()
root.title("Gerador de Documento com Usuário e Senha")
root.geometry("650x600")

tk.Label(root, text="Comprimento da Senha:").grid(row=0, column=0, padx=10, pady=5, sticky="w")
password_length_spinbox = tk.Spinbox(root, from_=8, to=20, width=5)
password_length_spinbox.grid(row=0, column=1, padx=10, pady=5)

generate_button = tk.Button(root, text="Gerar Senhas", width=20, height=2, command=generate_multiple_passwords)
generate_button.grid(row=1, column=0, columnspan=2, padx=10, pady=10)

tk.Label(root, text="Senhas Geradas").grid(row=2, column=0, padx=10, pady=5)
passwords_text = tk.Text(root, height=10, width=30)
passwords_text.grid(row=3, column=0, padx=10, pady=10)

# Título para o campo de usuários com a legenda em vermelho
tk.Label(root, text="Usuários").grid(row=2, column=1, padx=10, pady=5)
users_text = tk.Text(root, height=10, width=30)
users_text.grid(row=3, column=1, padx=10, pady=(0, 5))
users_text.bind("<KeyRelease>", on_users_text_change)
tk.Label(root, text="Cole aqui os usuários para serem mesclados", fg="red", font=("TkDefaultFont", 8)).grid(row=4, column=1, padx=10, sticky="n")

# Campo para exibir a mesclagem
tk.Label(root, text="Mesclagem de Usuários e Senhas").grid(row=5, column=0, columnspan=2, padx=10, pady=5)
merged_text = tk.Text(root, height=10, width=50)
merged_text.grid(row=6, column=0, columnspan=3, padx=10, pady=10)

doc_button = tk.Button(root, text="Gerar Documento", command=generate_doc)
doc_button.grid(row=7, column=0, columnspan=2, pady=20)

users = []
passwords = []

root.mainloop()
