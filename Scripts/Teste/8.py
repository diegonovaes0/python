import tkinter as tk
from tkinter import messagebox
from docx import Document
from docx.shared import Pt, RGBColor
from reportlab.pdfgen import canvas
import random
import string

# Função para gerar documento em docx
def create_vpn_doc(users_passwords, doc_name="Usuarios_Senhas_Details.docx"):
    doc = Document()
    
    # Conteúdo da capa
    doc.add_paragraph("LIBERAÇÃO PROJETO CLOUD AUTOSKY", style='Title')
    doc.add_paragraph("Este documento tem como objetivo apresentar as informações sobre arquitetura do projeto e as credenciais de acesso aos servidores.", style='Normal')

    # Tabela de informações
    for user, password in users_passwords:
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        
        # Linha 1: Usuário
        cell1 = table.cell(0, 0)
        run1 = cell1.paragraphs[0].add_run('Usuário')
        run1.font.name = 'Roboto'
        run1.font.size = Pt(12)
        run1.font.bold = True

        cell2 = table.cell(0, 1)
        run2 = cell2.paragraphs[0].add_run(user)
        run2.font.name = 'Calibri'
        run2.font.size = Pt(11)

        # Linha 2: Senha
        cell3 = table.cell(1, 0)
        run3 = cell3.paragraphs[0].add_run('Senha')
        run3.font.name = 'Roboto'
        run3.font.size = Pt(12)
        run3.font.bold = True

        cell4 = table.cell(1, 1)
        run4 = cell4.paragraphs[0].add_run(password if password else "(sem senha)")
        run4.font.name = 'Calibri'
        run4.font.size = Pt(11)
        
        doc.add_paragraph("")  # Espaço entre tabelas

    doc.save(doc_name)
    messagebox.showinfo("Sucesso", f"Documento DOCX salvo como {doc_name}")

# Função para exportar o mesmo conteúdo em PDF
def create_vpn_pdf(users_passwords, pdf_name="Usuarios_Senhas_Details.pdf"):
    c = canvas.Canvas(pdf_name)
    c.setFont("Helvetica-Bold", 14)
    c.drawString(100, 800, "LIBERAÇÃO PROJETO CLOUD AUTOSKY")
    
    # Subtítulo e descrição
    c.setFont("Helvetica", 10)
    text = c.beginText(50, 760)
    text.textLines("Este documento tem como objetivo apresentar as informações sobre arquitetura do projeto e as credenciais de acesso aos servidores.\n")
    c.drawText(text)

    # Dados em tabela
    y = 700
    for user, password in users_passwords:
        c.setFont("Helvetica-Bold", 10)
        c.drawString(50, y, f"Usuário: {user}")
        c.setFont("Helvetica", 10)
        c.drawString(200, y, f"Senha: {password}")
        y -= 20  # Nova linha

    c.save()
    messagebox.showinfo("Sucesso", f"Documento PDF salvo como {pdf_name}")

# Função chamada quando o botão é clicado
def generate_doc():
    users_passwords_text = merged_text.get("1.0", tk.END).strip()
    if not users_passwords_text:
        messagebox.showerror("Erro", "Por favor, insira os usuários e senhas.")
        return

    users_passwords = []
    lines = users_passwords_text.splitlines()
    for i in range(0, len(lines), 2):
        user = lines[i]
        password = lines[i + 1] if i + 1 < len(lines) else "(sem senha)"
        users_passwords.append((user, password))

    create_vpn_doc(users_passwords)  # Salvar DOCX
    create_vpn_pdf(users_passwords)  # Salvar PDF

# Interface gráfica
root = tk.Tk()
root.title("Exportar Documento")
root.geometry("500x400")

merged_text = tk.Text(root, height=15, width=50)
merged_text.pack(pady=20)

# Botões de Exportação
export_docx_button = tk.Button(root, text="Exportar como DOCX", command=lambda: generate_doc())
export_docx_button.pack(pady=10)

export_pdf_button = tk.Button(root, text="Exportar como PDF", command=lambda: generate_doc())
export_pdf_button.pack(pady=10)

root.mainloop()
