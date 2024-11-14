import os
import tkinter as tk
from tkinter import filedialog, messagebox
from PIL import Image
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from io import BytesIO

def export_to_docx():
    doc = Document()

    # Adicionar cabeçalho com imagem centralizada
    header = doc.sections[0].header
    img_path_header = os.path.join("img", "cabec_new.jpg")
    if os.path.exists(img_path_header):
        header_paragraph = header.paragraphs[0]
        header_paragraph.add_run().add_picture(img_path_header, width=Inches(6.3))  # Ajusta a largura da imagem para preencher
        header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Adicionar rodapé com imagem centralizada
    footer = doc.sections[0].footer
    img_path_footer = os.path.join("img", "rodape_new.jpg")
    if os.path.exists(img_path_footer):
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.add_run().add_picture(img_path_footer, width=Inches(6.3))
        footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Salvar o documento
    file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("DOCX files", "*.docx")])
    if file_path:
        doc.save(file_path)
        messagebox.showinfo("Exportação", f"Documento exportado com sucesso para {file_path}")

def export_to_pdf():
    pdf_path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF files", "*.pdf")])
    if pdf_path:
        pdf_canvas = canvas.Canvas(pdf_path, pagesize=A4)
        width, height = A4

        # Adicionar imagem de cabeçalho centralizada
        img_path_header = os.path.join("img", "cabec_new.jpg")
        if os.path.exists(img_path_header):
            img_header = Image.open(img_path_header)
            img_header.thumbnail((width, 100))  # Ajustar altura do cabeçalho conforme necessário
            img_io_header = BytesIO()
            img_header.save(img_io_header, format="PNG")
            img_reader_header = ImageReader(img_io_header)
            pdf_canvas.drawImage(img_reader_header, (width - img_header.width) / 2, height - 100)

        # Adicionar imagem de rodapé centralizada
        img_path_footer = os.path.join("img", "rodape_new.jpg")
        if os.path.exists(img_path_footer):
            img_footer = Image.open(img_path_footer)
            img_footer.thumbnail((width, 100))  # Ajustar altura do rodapé conforme necessário
            img_io_footer = BytesIO()
            img_footer.save(img_io_footer, format="PNG")
            img_reader_footer = ImageReader(img_io_footer)
            pdf_canvas.drawImage(img_reader_footer, (width - img_footer.width) / 2, 20)

        pdf_canvas.showPage()
        pdf_canvas.save()
        messagebox.showinfo("Exportação", f"Documento exportado com sucesso para {pdf_path}")

# Interface gráfica para os botões de exportação
root = tk.Tk()
root.title("Exportar Cabeçalho e Rodapé")
root.geometry("400x200")

btn_docx = tk.Button(root, text="Exportar para DOCX", command=export_to_docx, width=20)
btn_docx.pack(pady=20)

btn_pdf = tk.Button(root, text="Exportar para PDF", command=export_to_pdf, width=20)
btn_pdf.pack(pady=20)

root.mainloop()
