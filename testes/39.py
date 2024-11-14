import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from PIL import Image
from io import BytesIO

# Caminho das imagens
img_path_header = os.path.join("img", "cabec_new.jpg")
img_path_footer = os.path.join("img", "rodape_new.jpg")

# Função para criar o DOCX com cabeçalho e rodapé cobrindo toda a largura
def create_docx():
    doc = Document()

    # Ajuste das margens para maximizar o espaço disponível
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)

    # Adicionar imagem no cabeçalho cobrindo toda a largura
    if os.path.exists(img_path_header):
        header = doc.sections[0].header
        header_paragraph = header.paragraphs[0]
        header_paragraph.add_run().add_picture(img_path_header, width=Inches(8.27))  # Largura A4
        header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Adicionar imagem no rodapé cobrindo toda a largura
    if os.path.exists(img_path_footer):
        footer = doc.sections[0].footer
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.add_run().add_picture(img_path_footer, width=Inches(8.27))  # Largura A4
        footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Salvar o documento
    doc.save("output_document.docx")
    print("Documento DOCX criado com cabeçalho e rodapé em largura total.")

# Função para criar o PDF com cabeçalho e rodapé cobrindo toda a largura
def create_pdf():
    pdf_canvas = canvas.Canvas("output_document.pdf", pagesize=A4)
    width, height = A4

    # Adicionar imagem de cabeçalho cobrindo toda a largura
    if os.path.exists(img_path_header):
        img_header = Image.open(img_path_header)
        img_header = img_header.resize((int(width), 100))  # Ajuste a altura conforme necessário
        img_io_header = BytesIO()
        img_header.save(img_io_header, format="PNG")
        img_reader_header = ImageReader(img_io_header)
        pdf_canvas.drawImage(img_reader_header, 0, height - 100, width=width, height=100)

    # Adicionar imagem de rodapé cobrindo toda a largura
    if os.path.exists(img_path_footer):
        img_footer = Image.open(img_path_footer)
        img_footer = img_footer.resize((int(width), 100))  # Ajuste a altura conforme necessário
        img_io_footer = BytesIO()
        img_footer.save(img_io_footer, format="PNG")
        img_reader_footer = ImageReader(img_io_footer)
        pdf_canvas.drawImage(img_reader_footer, 0, 0, width=width, height=100)

    pdf_canvas.showPage()
    pdf_canvas.save()
    print("Documento PDF criado com cabeçalho e rodapé em largura total.")

# Executar as funções
create_docx()
create_pdf()
