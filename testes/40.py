from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_docx():
    # Criação do documento
    doc = Document()
    
    # Configuração inicial do documento
    title = "Relatório de Atividades"
    subtitle = "Resumo do Mês"
    paragraphs = [
        "Este documento apresenta um resumo das atividades realizadas no último mês.",
        "Foram abordados os principais pontos para a tomada de decisão estratégica da empresa."
    ]
    
    # Adicionando o título
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run(title)
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Adicionando o subtítulo
    subtitle_paragraph = doc.add_paragraph()
    subtitle_run = subtitle_paragraph.add_run(subtitle)
    subtitle_run.font.size = Pt(18)
    subtitle_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Adicionando parágrafos de texto
    for paragraph in paragraphs:
        para = doc.add_paragraph(paragraph)
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Função para adicionar texto colorido
    def add_colored_text(paragraph, text, color):
        run = paragraph.add_run(text)
        run.font.size = Pt(12)
        # Definindo a cor com RGBColor
        run.font.color.rgb = {'blue': RGBColor(0, 102, 204), 'black': RGBColor(0, 0, 0)}[color]
    
    # Exemplo de uso da função para adicionar texto colorido
    colored_paragraph = doc.add_paragraph()
    add_colored_text(colored_paragraph, "Este texto está em azul.", 'blue')
    add_colored_text(colored_paragraph, " E este texto está em preto.", 'black')
    
    # Salvando o documento
    doc.save("relatorio_atividades.docx")

# Executando a função para criar o documento
create_docx()
