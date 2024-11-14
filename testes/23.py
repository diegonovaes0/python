import requests
from docx import Document
from io import BytesIO
import re

def get_file_id_from_url(url):
    # Extrai o file_id da URL do Google Drive
    match = re.search(r'd/([a-zA-Z0-9_-]+)|id=([a-zA-Z0-9_-]+)', url)
    if match:
        return match.group(1) or match.group(2)
    else:
        print("URL inválida. Certifique-se de que é um link do Google Drive.")
        return None

def download_docx_from_google_docs(file_id):
    # URL de download para exportar como .docx
    download_url = f"https://docs.google.com/document/d/{file_id}/export?format=docx"
    
    # Baixa o arquivo
    response = requests.get(download_url, stream=True)
    if response.status_code == 200:
        print("Download concluído com sucesso.")
        return BytesIO(response.content)
    else:
        print("Falha ao baixar o arquivo. Verifique as permissões do documento.")
        return None

def read_docx_content(docx_file):
    # Carrega o documento do conteúdo baixado
    doc = Document(docx_file)
    content = {"paragraphs": [], "tables": []}

    # Extrai parágrafos
    for paragraph in doc.paragraphs:
        content["paragraphs"].append(paragraph.text)

    # Extrai tabelas
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        content["tables"].append(table_data)

    return content

# Solicita a URL ao usuário
url = input("Por favor, insira a URL do documento no Google Drive: ")
file_id = get_file_id_from_url(url)

# Baixa e processa o documento
if file_id:
    docx_file = download_docx_from_google_docs(file_id)
    if docx_file:
        content = read_docx_content(docx_file)

        # Exibe o conteúdo extraído
        print("Parágrafos:")
        for paragraph in content["paragraphs"]:
            print(paragraph)

        print("\nTabelas:")
        for table in content["tables"]:
            for row in table:
                print(row)
