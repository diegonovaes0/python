from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_docx():
    # Criação do documento
    doc = Document()
    
    # Adicionando o título
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run("Relatório de Atividades")
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Adicionando o subtítulo
    subtitle_paragraph = doc.add_paragraph()
    subtitle_run = subtitle_paragraph.add_run("Resumo do Mês")
    subtitle_run.font.size = Pt(18)
    subtitle_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Adicionando parágrafos de texto
    doc.add_paragraph("Este documento apresenta um resumo das atividades realizadas no último mês.")
    doc.add_paragraph("Foram abordados os principais pontos para a tomada de decisão estratégica da empresa.")

    # Adicionando parágrafo com texto colorido
    colored_paragraph = doc.add_paragraph()
    colored_text = colored_paragraph.add_run("Texto em azul.")
    colored_text.font.size = Pt(12)
    colored_text.font.color.rgb = RGBColor(0, 102, 204)

    # Salvando o documento
    doc.save("relatorio_atividades.docx")

# Executando a função para criar o documento
create_docx()
