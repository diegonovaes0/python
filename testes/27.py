import os
import tkinter as tk
from tkinter import filedialog, messagebox
from PIL import Image, ImageTk
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from io import BytesIO

def load_image(filename, width=None, height=None):
    img_path = os.path.join("img", filename)
    if os.path.exists(img_path):
        img = Image.open(img_path)
        if width and height:
            img = img.resize((width, height), Image.LANCZOS)
        return ImageTk.PhotoImage(img)
    else:
        print(f"Imagem {filename} não encontrada.")
        return None

def display_document_content():
    root = tk.Tk()
    root.title("Visualizador de Documento")
    root.geometry("800x900")

    # Canvas e Scrollbar para rolagem
    canvas = tk.Canvas(root, width=780, height=850)
    scrollbar = tk.Scrollbar(root, orient="vertical", command=canvas.yview)
    scrollable_frame = tk.Frame(canvas)
    scrollable_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
    canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
    canvas.configure(yscrollcommand=scrollbar.set)

    # Adiciona a capa
    capa_img = load_image("capa_doc_autosky.png", width=800, height=1132)
    if capa_img:
        capa_label = tk.Label(scrollable_frame, image=capa_img)
        capa_label.image = capa_img
        capa_label.pack(pady=10)
    
    # Adiciona o cabeçalho
    cabec_img = load_image("cabec_new.jpg", width=800, height=121)
    if cabec_img:
        cabec_label = tk.Label(scrollable_frame, image=cabec_img)
        cabec_label.image = cabec_img
        cabec_label.pack(pady=10)

    # Título centralizado
    title_label = tk.Label(scrollable_frame, text="LIBERAÇÃO PROJETO CLOUD AUTOSKY", font=("Helvetica", 14, "bold"))
    title_label.pack(pady=5)

    text_content = [
        "Este documento tem como objetivo apresentar as informações sobre arquitetura do projeto e as credenciais de acesso aos servidores.",
        "O ambiente Autosky é composto por um servidor de aplicação, onde sua principal finalidade é armazenar a instalação das aplicações client-server, e a virtualização para acesso dos usuários através do Autosky Platform.",
        "Demais finalidades como Banco de Dados, Serviços WEB, File Server, assim como outros serviços específicos, serão implantados conforme o modelo de arquitetura contemplado no escopo deste projeto, podendo ou não ocorrer a existência de mais servidores na arquitetura.",
        "Durante o período de implantação, treinamentos serão ministrados para adequar parceiro / cliente ao universo cloud Autosky, capacitando-os na aderência das melhores práticas e usabilidade dos recursos e serviços oferecidos pela Skyone.",
        "Os servidores são compostos por dois endereços de IP, sendo Público e Privado, assim como um endereço de DNS no qual poderá ser utilizado no lugar do IP público. Para acessar as instâncias via terminal server (RDP ou SSH), deverá ser utilizado o IP Público.",
        "No entanto, se a arquitetura deste projeto contempla serviços de VPN, após o fechamento e configuração das rotas, é possível realizar o acesso também pelo IP privado."
    ]
    
    for paragraph in text_content:
        text_label = tk.Label(scrollable_frame, text=paragraph, wraplength=750, justify="left", anchor="w", font=("Helvetica", 12))
        text_label.pack(pady=5, padx=5)

    # Rodapé
    rodape_img = load_image("rodape_new.jpg", width=800, height=102)
    if rodape_img:
        rodape_label = tk.Label(scrollable_frame, image=rodape_img)
        rodape_label.image = rodape_img
        rodape_label.pack(pady=10)

    canvas.pack(side="left", fill="both", expand=True)
    scrollbar.pack(side="right", fill="y")

    # Funções para exportar
    def export_to_docx():
        doc = Document()
        doc.add_paragraph("LIBERAÇÃO PROJETO CLOUD AUTOSKY", style="Title")
        
        for paragraph in text_content:
            doc.add_paragraph(paragraph)

        # Adiciona imagens na exportação para DOCX
        img_path = os.path.join("img", "capa_doc_autosky.png")
        if os.path.exists(img_path):
            doc.add_picture(img_path, width=docx.shared.Inches(6))

        # Salvamento com janela para escolher local
        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("DOCX files", "*.docx")])
        if file_path:
            doc.save(file_path)
            messagebox.showinfo("Exportação", f"Documento exportado com sucesso para {file_path}")

    def export_to_pdf():
        pdf_path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF files", "*.pdf")])
        if pdf_path:
            c = canvas.Canvas(pdf_path, pagesize=A4)
            width, height = A4
            y = height - 40
            c.setFont("Helvetica", 12)

            # Adiciona título
            c.drawString(40, y, "LIBERAÇÃO PROJETO CLOUD AUTOSKY")
            y -= 30

            for paragraph in text_content:
                c.drawString(40, y, paragraph)
                y -= 15
                if y < 40:
                    c.showPage()
                    y = height - 40

            # Adiciona a imagem ao PDF
            img_path = os.path.join("img", "capa_doc_autosky.png")
            if os.path.exists(img_path):
                img = Image.open(img_path)
                img.thumbnail((500, 700))  # Redimensiona para o PDF
                img_io = BytesIO()
                img.save(img_io, format="PNG")
                img_reader = ImageReader(img_io)
                c.drawImage(img_reader, 40, y - 300, width=400, height=300)
                y -= 320

            c.save()
            messagebox.showinfo("Exportação", f"Documento exportado com sucesso para {pdf_path}")

    # Botões para exportar
    btn_frame = tk.Frame(root)
    btn_frame.pack(pady=10)
    
    btn_docx = tk.Button(btn_frame, text="Exportar para DOCX", command=export_to_docx, width=20)
    btn_docx.grid(row=0, column=0, padx=5)
    
    btn_pdf = tk.Button(btn_frame, text="Exportar para PDF", command=export_to_pdf, width=20)
    btn_pdf.grid(row=0, column=1, padx=5)

    root.mainloop()

display_document_content()
