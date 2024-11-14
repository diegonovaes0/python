import requests
from docx import Document
from io import BytesIO
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
from PIL import Image, ImageTk
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
import re

def get_file_id_from_url(url):
    match = re.search(r'd/([a-zA-Z0-9_-]+)|id=([a-zA-Z0-9_-]+)', url)
    if match:
        return match.group(1) or match.group(2)
    else:
        messagebox.showerror("Erro", "URL inválida. Certifique-se de que é um link do Google Drive.")
        return None

def download_docx_from_google_docs(file_id):
    download_url = f"https://docs.google.com/document/d/{file_id}/export?format=docx"
    response = requests.get(download_url, stream=True)
    if response.status_code == 200:
        print("Download concluído com sucesso.")
        return BytesIO(response.content)
    else:
        messagebox.showerror("Erro", "Falha ao baixar o arquivo. Verifique as permissões do documento.")
        return None

def read_docx_content(docx_file):
    doc = Document(docx_file)
    content = {"paragraphs": [], "images": []}
    
    # Extrai parágrafos
    for paragraph in doc.paragraphs:
        content["paragraphs"].append(paragraph.text)
    
    # Extrai imagens
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            img_data = rel.target_part.blob
            content["images"].append(img_data)
    
    return content

def display_document_content(content):
    root = tk.Tk()
    root.title("Visualizador de Documento")
    root.geometry("800x600")

    text_area = scrolledtext.ScrolledText(root, wrap=tk.WORD, width=70, height=20, font=("Helvetica", 12))
    text_area.pack(pady=10)
    
    for paragraph in content["paragraphs"]:
        text_area.insert(tk.END, paragraph + "\n\n")

    for img_data in content["images"]:
        img = Image.open(BytesIO(img_data))
        img.thumbnail((400, 300))
        img_tk = ImageTk.PhotoImage(img)
        img_label = tk.Label(root, image=img_tk)
        img_label.image = img_tk
        img_label.pack(pady=5)

    def export_to_pdf():
        pdf_path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF files", "*.pdf")])
        if pdf_path:
            c = canvas.Canvas(pdf_path, pagesize=A4)
            width, height = A4
            y = height - 40
            c.setFont("Helvetica", 12)

            for paragraph in content["paragraphs"]:
                for line in paragraph.splitlines():
                    c.drawString(40, y, line)
                    y -= 14
                    if y < 40:
                        c.showPage()
                        y = height - 40
            
            for img_data in content["images"]:
                img = Image.open(BytesIO(img_data))
                img.thumbnail((400, 300))
                img_path = "/tmp/temp_img.png"
                img.save(img_path)
                if y < 100:
                    c.showPage()
                    y = height - 100
                c.drawImage(img_path, 40, y - 300)
                y -= 320

            c.save()
            messagebox.showinfo("Sucesso", f"Documento salvo como PDF em {pdf_path}")

    def export_to_docx():
        docx_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("DOCX files", "*.docx")])
        if docx_path:
            new_doc = Document()
            for paragraph in content["paragraphs"]:
                new_doc.add_paragraph(paragraph)

            for img_data in content["images"]:
                img = Image.open(BytesIO(img_data))
                img_path = "/tmp/temp_img_export.png"
                img.save(img_path)
                new_doc.add_picture(img_path, width=Image.open(img_path).width * 0.5)

            new_doc.save(docx_path)
            messagebox.showinfo("Sucesso", f"Documento salvo como DOCX em {docx_path}")

    btn_pdf = tk.Button(root, text="Exportar para PDF", command=export_to_pdf)
    btn_pdf.pack(pady=5)
    
    btn_docx = tk.Button(root, text="Exportar para DOCX", command=export_to_docx)
    btn_docx.pack(pady=5)

    root.mainloop()

url = input("Por favor, insira a URL do documento no Google Drive: ")
file_id = get_file_id_from_url(url)

if file_id:
    docx_file = download_docx_from_google_docs(file_id)
    if docx_file:
        content = read_docx_content(docx_file)
        display_document_content(content)
