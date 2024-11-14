import os
import tkinter as tk
from tkinter import filedialog, messagebox
from PIL import Image, ImageTk
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from io import BytesIO

def load_image(filename, width=None, height=None):
    img_path = os.path.join("img", filename)
    if os.path.exists(img_path):
        img = Image.open(img_path)
        if width and height:
            img = img.resize((width, height), Image.LANCZOS)
        return ImageTk.PhotoImage(img)
    else:
        print(f"Imagem {filename} não encontrada.")
        return None

def display_document_content():
    root = tk.Tk()
    root.title("Visualizador de Documento")
    root.geometry("800x900")
    root.configure(bg="white")

    main_frame = tk.Frame(root, bg="white")
    main_frame.pack(fill="both", expand=True)

    doc_canvas = tk.Canvas(main_frame, bg="white", width=780, height=850)
    scrollbar = tk.Scrollbar(main_frame, orient="vertical", command=doc_canvas.yview)
    scrollable_frame = tk.Frame(doc_canvas, bg="white")

    scrollable_frame.bind("<Configure>", lambda e: doc_canvas.configure(scrollregion=doc_canvas.bbox("all")))
    doc_canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
    doc_canvas.configure(yscrollcommand=scrollbar.set)

    capa_img = load_image("capa_doc_autosky.png", width=800, height=1132)
    if capa_img:
        capa_label = tk.Label(scrollable_frame, image=capa_img, bg="white")
        capa_label.image = capa_img
        capa_label.pack(pady=10)
    
    cabec_img = load_image("cabec_new.jpg", width=800, height=121)
    if cabec_img:
        cabec_label = tk.Label(scrollable_frame, image=cabec_img, bg="white")
        cabec_label.image = cabec_img
        cabec_label.pack(pady=10)

    title_label = tk.Label(scrollable_frame, text="LIBERAÇÃO PROJETO CLOUD AUTOSKY", font=("Helvetica", 14, "bold"), bg="white")
    title_label.pack(pady=5)

    text_content = [
        ("Este documento tem como objetivo apresentar as informações sobre arquitetura do projeto e as credenciais de acesso aos servidores.", False, False),
        ("O ambiente Autosky é composto por um servidor de aplicação, onde sua principal finalidade é armazenar a instalação das aplicações client-server, e a virtualização para acesso dos usuários através do Autosky Platform.", False, False),
        ("Demais finalidades como Banco de Dados, Serviços WEB, File Server, assim como outros serviços específicos, serão implantados conforme o modelo de arquitetura contemplado no escopo deste projeto, podendo ou não ocorrer a existência de mais servidores na arquitetura.", False, False),
        ("Durante o período de implantação, treinamentos serão ministrados para adequar parceiro / cliente ao universo cloud Autosky, capacitando-os na aderência das melhores práticas e usabilidade dos recursos e serviços oferecidos pela Skyone.", False, False),
        ("Os servidores são compostos por dois endereços de IP, sendo Público e Privado, assim como um endereço de DNS no qual poderá ser utilizado no lugar do IP público.", True, False),
        ("Para acessar as instâncias via terminal server (RDP ou SSH), deverá ser utilizado o IP Público.", False, True)
    ]
    
    for paragraph, bold, italic in text_content:
        text_label = tk.Label(scrollable_frame, text=paragraph, wraplength=750, justify="left", anchor="w", font=("Helvetica", 12, "bold" if bold else "italic" if italic else "normal"), bg="white")
        text_label.pack(pady=5, padx=5)

    rodape_img = load_image("rodape_new.jpg", width=800, height=102)
    if rodape_img:
        rodape_label = tk.Label(scrollable_frame, image=rodape_img, bg="white")
        rodape_label.image = rodape_img
        rodape_label.pack(pady=10)

    def export_to_docx():
        doc = Document()
        title = doc.add_paragraph("LIBERAÇÃO PROJETO CLOUD AUTOSKY")
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title.runs[0].bold = True

        for paragraph, bold, italic in text_content:
            p = doc.add_paragraph(paragraph)
            run = p.runs[0]
            run.bold = bold
            run.italic = italic
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # Adiciona uma quebra de página
        doc.add_page_break()

        # Adiciona a imagem ao documento
        img_path = os.path.join("img", "capa_doc_autosky.png")
        if os.path.exists(img_path):
            doc.add_picture(img_path, width=Inches(6))
            doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("DOCX files", "*.docx")])
        if file_path:
            doc.save(file_path)
            messagebox.showinfo("Exportação", f"Documento exportado com sucesso para {file_path}")

    def export_to_pdf():
        pdf_path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF files", "*.pdf")])
        if pdf_path:
            pdf_canvas = canvas.Canvas(pdf_path, pagesize=A4)
            width, height = A4
            y = height - 50
            pdf_canvas.setFont("Helvetica-Bold", 14)
            pdf_canvas.drawCentredString(width / 2, y, "LIBERAÇÃO PROJETO CLOUD AUTOSKY")
            y -= 40
            pdf_canvas.setFont("Helvetica", 12)

            for paragraph, bold, italic in text_content:
                if bold:
                    pdf_canvas.setFont("Helvetica-Bold", 12)
                elif italic:
                    pdf_canvas.setFont("Helvetica-Oblique", 12)
                else:
                    pdf_canvas.setFont("Helvetica", 12)
                
                pdf_canvas.drawString(40, y, paragraph)
                y -= 20

                if y < 50:
                    pdf_canvas.showPage()
                    y = height - 50

            # Adiciona imagem ao PDF com controle de páginas
            img_path = os.path.join("img", "capa_doc_autosky.png")
            if os.path.exists(img_path):
                img = Image.open(img_path)
                img.thumbnail((500, 700))
                img_io = BytesIO()
                img.save(img_io, format="PNG")
                img_reader = ImageReader(img_io)
                pdf_canvas.drawImage(img_reader, 40, y - 300, width=400, height=300)
                y -= 320

            pdf_canvas.save()
            messagebox.showinfo("Exportação", f"Documento exportado com sucesso para {pdf_path}")

    btn_frame = tk.Frame(root, bg="white")
    btn_frame.pack(pady=10)

    btn_docx = tk.Button(btn_frame, text="Exportar para DOCX", command=export_to_docx, width=20)
    btn_docx.grid(row=0, column=0, padx=5)
    
    btn_pdf = tk.Button(btn_frame, text="Exportar para PDF", command=export_to_pdf, width=20)
    btn_pdf.grid(row=0, column=1, padx=5)

    doc_canvas.pack(side="left", fill="both", expand=True)
    scrollbar.pack(side="right", fill="y")

    root.mainloop()

display_document_content()
